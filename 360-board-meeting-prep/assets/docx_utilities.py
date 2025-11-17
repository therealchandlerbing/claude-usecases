"""
DOCX Generation Utilities for 360 Board Meeting Prep
Provides formatted document generation following 360's exact specifications
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime


class BoardPacketFormatter:
    """Handles all DOCX formatting for 360 board packets"""

    # Typography specifications
    FONTS = {
        'header': 'Arial',
        'body': 'Calibri'
    }

    SIZES = {
        'title': Pt(24),
        'subtitle': Pt(13),
        'h1': Pt(14),
        'h2': Pt(12),
        'h3': Pt(11),
        'body': Pt(11)
    }

    SPACING = {
        'h1_before': Pt(18),
        'h1_after': Pt(6),
        'h2_before': Pt(12),
        'h2_after': Pt(6),
        'h3_before': Pt(6),
        'h3_after': Pt(6),
        'section_divider_before': Pt(18),
        'section_divider_after': Pt(18)
    }

    COLORS = {
        'table_label_bg': 'ECF0F1',  # Light gray for table labels
        'black': '000000'
    }

    def __init__(self):
        self.doc = Document()
        self._configure_document()

    def _configure_document(self):
        """Set up document-level defaults"""
        # Set page margins (1 inch all sides)
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        # Configure styles
        self._configure_styles()

    def _configure_styles(self):
        """Configure document styles"""
        styles = self.doc.styles

        # Title style
        if 'Title' in styles:
            title_style = styles['Title']
            title_font = title_style.font
            title_font.name = self.FONTS['header']
            title_font.size = self.SIZES['title']
            title_font.bold = True

        # Heading 1 style
        if 'Heading 1' in styles:
            h1_style = styles['Heading 1']
            h1_font = h1_style.font
            h1_font.name = self.FONTS['header']
            h1_font.size = self.SIZES['h1']
            h1_font.bold = True
            h1_para = h1_style.paragraph_format
            h1_para.space_before = self.SPACING['h1_before']
            h1_para.space_after = self.SPACING['h1_after']

        # Heading 2 style
        if 'Heading 2' in styles:
            h2_style = styles['Heading 2']
            h2_font = h2_style.font
            h2_font.name = self.FONTS['header']
            h2_font.size = self.SIZES['h2']
            h2_font.bold = True
            h2_para = h2_style.paragraph_format
            h2_para.space_before = self.SPACING['h2_before']
            h2_para.space_after = self.SPACING['h2_after']

        # Heading 3 style
        if 'Heading 3' in styles:
            h3_style = styles['Heading 3']
            h3_font = h3_style.font
            h3_font.name = self.FONTS['header']
            h3_font.size = self.SIZES['h3']
            h3_font.bold = True
            h3_para = h3_style.paragraph_format
            h3_para.space_before = self.SPACING['h3_before']
            h3_para.space_after = self.SPACING['h3_after']

        # Normal/Body style
        if 'Normal' in styles:
            normal_style = styles['Normal']
            normal_font = normal_style.font
            normal_font.name = self.FONTS['body']
            normal_font.size = self.SIZES['body']

    @staticmethod
    def set_cell_background(cell, color_hex):
        """Set cell background color

        Args:
            cell: Table cell object
            color_hex: Hex color code (e.g., 'ECF0F1')
        """
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color_hex)
        cell._element.get_or_add_tcPr().append(shading_elm)

    @staticmethod
    def remove_table_borders(table):
        """Remove borders from a table

        Args:
            table: Table object
        """
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            tblBorders.append(border)

        tblPr.append(tblBorders)

    def create_meeting_info_table(self, meeting_info):
        """Create meeting information table (4x2, gray labels)

        Args:
            meeting_info: Dict with keys: date, time, location, meeting_type

        Returns:
            Table object
        """
        table = self.doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'

        # Define content
        info = [
            ("Date:", meeting_info['date']),
            ("Time:", meeting_info['time']),
            ("Location:", meeting_info['location']),
            ("Meeting Type:", meeting_info['meeting_type'])
        ]

        # Populate and format
        for i, (label, content) in enumerate(info):
            row = table.rows[i]

            # Left cell (label - gray background, bold)
            left_cell = row.cells[0]
            left_cell.text = label
            self.set_cell_background(left_cell, self.COLORS['table_label_bg'])
            left_cell.paragraphs[0].runs[0].font.bold = True
            left_cell.paragraphs[0].runs[0].font.name = self.FONTS['body']
            left_cell.paragraphs[0].runs[0].font.size = self.SIZES['body']

            # Right cell (content - white background)
            right_cell = row.cells[1]
            right_cell.text = content
            right_cell.paragraphs[0].runs[0].font.name = self.FONTS['body']
            right_cell.paragraphs[0].runs[0].font.size = self.SIZES['body']

        # Set column widths
        table.columns[0].width = Inches(1.5)  # ~25%
        table.columns[1].width = Inches(4.5)  # ~75%

        # Remove borders
        self.remove_table_borders(table)

        return table

    def create_motion_table(self, motion_num, title, motion_text):
        """Create motion tracking table (2x1, bordered)

        Args:
            motion_num: Motion number (int)
            title: Motion title (str)
            motion_text: Full motion text (str)

        Returns:
            Table object
        """
        # Add spacing before table
        self.doc.add_paragraph()

        # Create 2x1 table with border
        table = self.doc.add_table(rows=2, cols=1)
        table.style = 'Table Grid'

        # Row 1: Motion title (bold)
        title_cell = table.rows[0].cells[0]
        title_para = title_cell.paragraphs[0]
        title_run = title_para.add_run(f'MOTION {motion_num}: {title}')
        title_run.font.bold = True
        title_run.font.name = self.FONTS['body']
        title_run.font.size = self.SIZES['body']

        # Row 2: Motion details
        details_cell = table.rows[1].cells[0]

        # Motion text (italic, quoted)
        motion_para = details_cell.paragraphs[0]
        motion_run = motion_para.add_run(f'"{motion_text}"')
        motion_run.font.italic = True
        motion_run.font.name = self.FONTS['body']
        motion_run.font.size = self.SIZES['body']

        # Add line break and fields
        details_cell.add_paragraph()  # Spacing
        moved_para = details_cell.add_paragraph('Moved by: ______________')
        moved_para.runs[0].font.name = self.FONTS['body']
        moved_para.runs[0].font.size = self.SIZES['body']

        seconded_para = details_cell.add_paragraph('Seconded by: ______________')
        seconded_para.runs[0].font.name = self.FONTS['body']
        seconded_para.runs[0].font.size = self.SIZES['body']

        vote_para = details_cell.add_paragraph()
        vote_run1 = vote_para.add_run('Vote: ')
        vote_run1.font.name = self.FONTS['body']
        vote_run1.font.size = self.SIZES['body']

        vote_run2 = vote_para.add_run('☐ Approved  ☐ Denied  ☐ Abstained')
        vote_run2.font.name = self.FONTS['body']
        vote_run2.font.size = self.SIZES['body']

        return table

    def add_section_divider(self):
        """Add centered bullet divider (• • •)"""
        para = self.doc.add_paragraph('• • •')
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_format = para.paragraph_format
        para_format.space_before = self.SPACING['section_divider_before']
        para_format.space_after = self.SPACING['section_divider_after']

        # Format the divider
        para.runs[0].font.name = self.FONTS['body']
        para.runs[0].font.size = self.SIZES['body']

    def add_title(self, text, centered=True):
        """Add title text (24pt Arial Bold)

        Args:
            text: Title text
            centered: Whether to center (default True)
        """
        para = self.doc.add_paragraph(text)
        para.style = 'Title'
        if centered:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_subtitle(self, text, centered=True):
        """Add subtitle text (13pt Arial)

        Args:
            text: Subtitle text
            centered: Whether to center (default True)
        """
        para = self.doc.add_paragraph(text)
        para.style = 'Subtitle'
        if centered:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_heading1(self, text):
        """Add Heading 1 (14pt Arial Bold, uppercase)

        Args:
            text: Heading text
        """
        para = self.doc.add_paragraph(text.upper())
        para.style = 'Heading 1'

    def add_heading2(self, text):
        """Add Heading 2 (12pt Arial Bold)

        Args:
            text: Heading text
        """
        para = self.doc.add_paragraph(text)
        para.style = 'Heading 2'

    def add_heading3(self, text):
        """Add Heading 3 (11pt Arial Bold)

        Args:
            text: Heading text
        """
        para = self.doc.add_paragraph(text)
        para.style = 'Heading 3'

    def add_body_text(self, text):
        """Add body paragraph (11pt Calibri)

        Args:
            text: Body text
        """
        para = self.doc.add_paragraph(text)
        para.style = 'Normal'

    def add_bullet_list(self, items):
        """Add bulleted list

        Args:
            items: List of strings
        """
        for item in items:
            para = self.doc.add_paragraph(item, style='List Bullet')

    def save(self, filename):
        """Save document to file

        Args:
            filename: Output filename
        """
        self.doc.save(filename)
        return filename


def create_board_packet(meeting_info, motions=None, include_motion_tracking=True):
    """Main function to create complete board packet

    Args:
        meeting_info: Dict with meeting details
        motions: List of motion dicts (optional)
        include_motion_tracking: Whether to include motion tracking sheet

    Returns:
        BoardPacketFormatter object with complete document
    """
    formatter = BoardPacketFormatter()

    # Cover page
    formatter.add_title('360 SOCIAL IMPACT STUDIOS')
    formatter.add_subtitle(meeting_info['meeting_type'])
    formatter.add_subtitle(meeting_info['full_date'])

    # Meeting information table
    formatter.doc.add_paragraph()  # Spacing
    formatter.create_meeting_info_table(meeting_info)

    # Motion tracking (if requested and motions provided)
    if include_motion_tracking and motions:
        formatter.doc.add_paragraph()  # Spacing
        formatter.add_section_divider()

        formatter.add_heading1('MOTION TRACKING SHEET')

        # Instructions
        instructions = formatter.doc.add_paragraph()
        instructions.add_run('This sheet contains pre-drafted motions for board consideration. ')
        instructions.add_run('During the meeting:\n\n')
        instructions.add_run('1. The Chair will call for a motion using the language provided\n')
        instructions.add_run('2. A board member moves (their name recorded in "Moved by")\n')
        instructions.add_run('3. Another board member seconds (their name recorded in "Seconded by")\n')
        instructions.add_run('4. Discussion occurs if needed\n')
        instructions.add_run('5. The Chair calls for a vote\n')
        instructions.add_run('6. The Secretary marks the appropriate box and records the vote')

        # Add motions
        for motion in motions:
            formatter.create_motion_table(
                motion['number'],
                motion['title'],
                motion['text']
            )

    return formatter


# Example usage
if __name__ == "__main__":
    # Sample data
    meeting_info = {
        'meeting_type': 'Annual Board Meeting',
        'full_date': 'Monday, November 17, 2025',
        'date': 'Monday, November 17, 2025',
        'time': '1:00-2:30 PM PST',
        'location': 'Online (Zoom)',
    }

    sample_motions = [
        {
            'number': 1,
            'title': 'Approve Meeting Agenda',
            'text': 'I move that we approve the agenda for the Annual Board Meeting of November 17, 2025'
        },
        {
            'number': 13,
            'title': 'Adjournment',
            'text': 'I move to adjourn this meeting'
        }
    ]

    # Create board packet
    packet = create_board_packet(meeting_info, sample_motions)
    packet.save('/mnt/user-data/outputs/sample_board_packet.docx')

    print("Sample board packet created: /mnt/user-data/outputs/sample_board_packet.docx")
